import os
import re
from pathlib import Path
from typing import cast

import docx.document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.styles.style import ParagraphStyle
from docx.text.paragraph import Paragraph


def add_custom_element(caption: Paragraph, name: str, key: str, value: str):
    run = caption.add_run()
    element = OxmlElement(name)
    element.set(qn(key), value)
    run._r.append(element)

    return element


def replace_placeholder_with_figure(doc: docx.document.Document, base: Path, before: str, field_name: str):
    for p in doc.paragraphs:
        match = re.match(r"^\{\{ (.+?) \}\}$", p.text)
        if not match:
            continue

        p.text = ""  # Clear the placeholder text
        img_path = f"{match.group(1)}.png"  # Extract the image path from the placeholder

        width = 16 / 2.54  # Convert cm to inches, full width in default document
        p.add_run().add_picture(str(base / img_path), width=Inches(width))

        caption_p = p.insert_paragraph_before(style="Caption")
        p._p.addnext(caption_p._p)
        caption_p.add_run(f"{before} ")

        # 1. Begin marker
        add_custom_element(caption_p, "w:fldChar", "w:fldCharType", "begin")

        # 2. Field instruction (SEQ)
        txt = add_custom_element(caption_p, "w:instrText", "xml:space", "preserve")
        txt.text = f" SEQ {field_name} \\* ARABIC "

        # 3. Separate marker
        add_custom_element(caption_p, "w:fldChar", "w:fldCharType", "separate")

        # 4. Cached result (placeholder before update)
        caption_p.add_run("<invalid, update fields>")

        # 5. End marker
        add_custom_element(caption_p, "w:fldChar", "w:fldCharType", "end")

        # Add the actual description
        caption_p.add_run(": " + match.group(1))


def add_style(doc: docx.document.Document):
    styles = doc.styles

    if "Caption" in styles:
        return

    caption_style = styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)

    caption_style = cast(ParagraphStyle, caption_style)  # typehint for mypy

    caption_style.base_style = styles["Normal"]
    caption_style.font.name = "Atpos"
    caption_style.font.size = Pt(9)
    caption_style.font.italic = True

    caption_style.hidden = False
    caption_style.quick_style = True


def example():
    document = docx.Document("base.docx")
    add_style(document)

    replace_placeholder_with_figure(document, Path(os.getcwd()), "Figure", "Figure")

    document.save("edited.docx")


if __name__ == "__main__":
    example()
